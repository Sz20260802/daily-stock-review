#!/usr/bin/env python3
"""将每日复盘 JSON 导出为 Word(.docx) 和 HTML 文件
用法:
  python scripts/export_report.py                    # 导出最新
  python scripts/export_report.py --date 2026-08-11  # 导出指定日期
输出: C:\daily-stock-review\export\复盘报告_日期.docx 和 .html
"""
import argparse
import json
import sys
import webbrowser
from datetime import datetime
from pathlib import Path

BASE = Path(__file__).resolve().parent.parent
REVIEW_DIR = BASE / "data" / "reviews"
OUT_DIR = BASE / "export"
OUT_DIR.mkdir(exist_ok=True)

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("[warn] 未安装 python-docx，仅生成 HTML。安装: pip install python-docx")


def esc(s):
    return str(s if s is not None else "").replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")


def load_review(date):
    path = REVIEW_DIR / ("latest.json" if date == "latest" else f"{date}.json")
    if not path.exists():
        return None
    return json.loads(path.read_text(encoding="utf-8"))


# ================== Word 导出 ==================
def set_cjk(run, size=11, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._element.get_or_add_rPr()
    rFonts = r.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = r.makeelement(qn('w:rFonts'), {})
        r.append(rFonts)
    rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')


def add_para(doc, text, size=11, bold=False, color=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_cjk(run, size, bold, color)
    return p


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(str(h))
        set_cjk(run, 10, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for j, v in enumerate(row):
            cells[j].text = ""
            run = cells[j].paragraphs[0].add_run(str(v))
            set_cjk(run, 10)
    return t


def build_docx(data, date_str, out_path):
    doc = Document()
    meta = data.get("meta") or {}
    m = data.get("market") or {}

    add_para(doc, "📈 每日股票复盘报告", 18, True, (31, 56, 100), 4)
    add_para(doc, f"复盘日期：{meta.get('date','')}　生成时间：{meta.get('generated_at','')}", 9, False, (107,114,128), 12)

    add_para(doc, "一、市场总览", 14, True, (31,56,100), 8)
    indices = m.get("indices") or []
    if indices:
        add_table(doc, ["代码","名称","收盘点位","涨跌幅"],
                  [[x.get("code",""), x.get("name",""), x.get("value",""), f"{x.get('change_pct','')}%"] for x in indices])
    stt = m.get("stats") or {}
    stats_txt = "　".join(f"{k}：{v}" for k,v in [
        ("上涨", stt.get("advancers")), ("下跌", stt.get("decliners")), ("平盘", stt.get("flat")),
        ("涨停", stt.get("limit_up")), ("跌停", stt.get("limit_down"))] if v is not None)
    if stats_txt:
        add_para(doc, stats_txt, 10)
    if m.get("turnover_yi"):
        add_para(doc, f"两市成交：{m['turnover_yi']} 亿", 10)
    if m.get("broken_rate") is not None:
        add_para(doc, f"炸板率：{m['broken_rate']}%", 10)
    if m.get("summary"):
        add_para(doc, "📝 AI 今日小结：", 11, True, (37,99,235))
        add_para(doc, m["summary"], 11)

    add_para(doc, "二、涨停梯队与龙虎榜", 14, True, (31,56,100), 8)
    ladder = m.get("limit_up_ladder") or {}
    if ladder.get("max_board"):
        add_para(doc, f"最高板：{ladder['max_board']} 板", 10)
    if ladder.get("ladders"):
        for k, v in ladder["ladders"].items():
            add_para(doc, f"{k}：{'、'.join(v)}", 10)
    lhb = m.get("lhb") or {}
    if lhb.get("top10"):
        add_table(doc, ["名称","涨跌幅","净买额(亿)","上榜原因"],
                  [[x.get("name",""), f"{x.get('close_pct','')}%", x.get("net_buy_yi",""), x.get("reason","")] for x in lhb["top10"]])

    themes = data.get("themes") or []
    if themes:
        add_para(doc, "三、主题复盘", 14, True, (31,56,100), 8)
        for i, t in enumerate(themes, 1):
            h = t.get("header") or {}
            add_para(doc, f"主题 #{h.get('index', i)}　{h.get('title', t.get('title',''))}　[{h.get('confidence','')}]", 12, True, (30,64,175))
            if t.get("tags"):
                add_para(doc, "标签：" + "、".join(t["tags"]), 9, False, (107,114,128))
            radar = t.get("radar_scores") or {}
            if radar:
                add_para(doc, f"评分（供需/估值/阶段）：{radar.get('supply_demand','-')}/{radar.get('valuation','-')}/{radar.get('stage','-')}　总分 {radar.get('total','-')}　平均 {radar.get('average','-')}", 10)
            add_para(doc, f"核心观点：{t.get('core_thesis','')}", 10)
            inst = t.get("institutions") or {}
            if inst.get("list"):
                names = "、".join(f"{x.get('name','')}({x.get('report_id','')})" for x in inst["list"])
                add_para(doc, f"机构覆盖（{inst.get('count','')}家）：{names}", 9, False, (107,114,128))
            if t.get("evidence_table"):
                add_table(doc, ["指标","数值","来源"],
                          [[x.get("indicator",""), x.get("value",""), x.get("source_type","")] for x in t["evidence_table"]])
            if t.get("logic_chain"):
                add_para(doc, "逻辑链条：", 10, True)
                add_para(doc, t["logic_chain"], 10)
            if t.get("risk_note"):
                add_para(doc, f"⚠️ 风险提示：{t['risk_note']}", 10, False, (153,27,27))
            if t.get("falsifiability_threshold"):
                add_para(doc, f"🔍 可证伪阈值：{t['falsifiability_threshold']}", 10, False, (146,64,14))

    evs = data.get("calendar_events") or []
    if evs:
        add_para(doc, "四、事件日历", 14, True, (31,56,100), 8)
        for ev in evs:
            add_para(doc, f"{ev.get('date','')}　{ev.get('event_name','')}", 11, True)
            add_para(doc, f"影响：{ev.get('impact','')}", 10)
            if ev.get("data_verification_point"):
                add_para(doc, f"验证点：{ev.get('data_verification_point','')}", 10, False, (146,64,14))

    corrs = data.get("correlations") or []
    if corrs:
        add_para(doc, "五、跨主题关联图谱", 14, True, (31,56,100), 8)
        for c in corrs:
            add_para(doc, f"{c.get('from','')} → {c.get('to','')}（{c.get('relation','')}）", 11, True)
            add_para(doc, c.get("evidence",""), 10)

    its = data.get("insights") or []
    if its:
        add_para(doc, "六、关键洞察", 14, True, (31,56,100), 8)
        for i, it in enumerate(its, 1):
            if isinstance(it, str):
                add_para(doc, f"{i}. {it}", 10)
            else:
                add_para(doc, f"{i}. {it.get('title','')}", 11, True)
                add_para(doc, it.get("detail",""), 10)

    add_para(doc, "", 8)
    add_para(doc, "⚠️ 本报告仅供学习研究，不构成投资建议。数据来源：东方财富等公开渠道。", 9, False, (107,114,128))
    doc.save(str(out_path))
    return out_path


# ================== HTML 导出 ==================
def build_html(data, date_str, out_path):
    meta = data.get("meta") or {}
    m = data.get("market") or {}
    h = []
    h.append("<!DOCTYPE html><html lang='zh-CN'><head><meta charset='UTF-8'><title>每日股票复盘</title>")
    h.append("<style>body{font-family:'Microsoft YaHei',sans-serif;background:#f4f6f9;color:#1f2937;max-width:900px;margin:20px auto;padding:0 16px}")
    h.append("h1{color:#1e3a8a}.card{background:#fff;border:1px solid #e5e7eb;border-radius:12px;padding:16px;margin:12px 0}")
    h.append(".sec{color:#1e3a8a;border-left:4px solid #2563eb;padding-left:10px;margin:18px 0 10px}")
    h.append("table{width:100%;border-collapse:collapse;font-size:13px;background:#fff}")
    h.append("th{background:#f1f5f9;text-align:left;padding:6px 8px}td{border-top:1px solid #e5e7eb;padding:6px 8px}")
    h.append(".chip{display:inline-block;background:#eef2ff;color:#3730a3;border-radius:999px;padding:2px 10px;font-size:12px;margin:2px 4px 2px 0}")
    h.append(".up{color:#dc2626}.down{color:#16a34a}.muted{color:#6b7280;font-size:12px}")
    h.append(".thesis{background:#f8fafc;border-left:3px solid #2563eb;padding:10px 12px;border-radius:8px;font-size:14px}")
    h.append(".risk{background:#fef2f2;border:1px solid #fecaca;color:#991b1b;border-radius:8px;padding:8px 12px;font-size:13px;margin-top:8px}")
    h.append(".falsify{background:#fffbeb;border:1px solid #fde68a;color:#92400e;border-radius:8px;padding:8px 12px;font-size:13px;margin-top:8px}</style></head><body>")

    h.append("<h1>📈 每日股票复盘报告</h1>")
    h.append(f"<p class='muted'>复盘日期：{esc(meta.get('date',''))}　生成时间：{esc(meta.get('generated_at',''))}</p>")

    h.append("<div class='sec'><b>一、市场总览</b></div>")
    indices = m.get("indices") or []
    if indices:
        h.append("<table><tr><th>名称</th><th>收盘点位</th><th>涨跌幅</th></tr>")
        for x in indices:
            cls = "up" if (x.get("change_pct") or 0) >= 0 else "down"
            sign = "+" if (x.get("change_pct") or 0) >= 0 else ""
            h.append(f"<tr><td>{esc(x.get('name',''))}</td><td>{x.get('value','-')}</td><td class='{cls}'>{sign}{x.get('change_pct','-')}%</td></tr>")
        h.append("</table>")
    stt = m.get("stats") or {}
    chips = []
    for k, label in [("advancers","上涨"),("decliners","下跌"),("flat","平盘"),("limit_up","涨停"),("limit_down","跌停")]:
        if stt.get(k) is not None:
            chips.append(f"<span class='chip'>{label} {stt[k]}</span>")
    if m.get("turnover_yi"):
        chips.append(f"<span class='chip'>两市成交 {m['turnover_yi']} 亿</span>")
    if m.get("broken_rate") is not None:
        chips.append(f"<span class='chip'>炸板率 {m['broken_rate']}%</span>")
    if chips:
        h.append("<div>" + "".join(chips) + "</div>")
    if m.get("summary"):
        h.append(f"<div class='thesis' style='margin-top:10px'><b>📝 AI 今日小结：</b><br>{esc(m['summary'])}</div>")

    h.append("<div class='sec'><b>二、涨停梯队与龙虎榜</b></div>")
    ladder = m.get("limit_up_ladder") or {}
    if ladder.get("max_board"):
        h.append(f"<span class='chip'>最高板：{ladder['max_board']} 板</span>")
    if ladder.get("ladders"):
        h.append("<div class='card'>")
        for k, v in ladder["ladders"].items():
            h.append(f"<span class='chip' style='background:#1e3a8a;color:#fff'>{esc(k)}</span> {esc('、'.join(v))}<br>")
        h.append("</div>")
    lhb = m.get("lhb") or {}
    if lhb.get("top10"):
        h.append("<table><tr><th>名称</th><th>涨跌幅</th><th>净买额(亿)</th><th>上榜原因</th></tr>")
        for x in lhb["top10"]:
            cls = "up" if (x.get("close_pct") or 0) >= 0 else "down"
            sign = "+" if (x.get("close_pct") or 0) >= 0 else ""
            h.append(f"<tr><td>{esc(x.get('name',''))}</td><td class='{cls}'>{sign}{x.get('close_pct','-')}%</td><td>{x.get('net_buy_yi','-')}</td><td>{esc(x.get('reason',''))}</td></tr>")
        h.append("</table>")

    themes = data.get("themes") or []
    if themes:
        h.append("<div class='sec'><b>三、主题复盘</b></div>")
        for i, t in enumerate(themes, 1):
            hd = t.get("header") or {}
            h.append("<div class='card'><b>")
            h.append(f"主题 #{hd.get('index',i)}　{esc(hd.get('title',t.get('title','')))}　[{esc(hd.get('confidence',''))}]")
            h.append("</b>")
            if t.get("tags"):
                h.append("<div>" + "".join(f"<span class='chip'>{esc(x)}</span>" for x in t["tags"]) + "</div>")
            radar = t.get("radar_scores") or {}
            if radar:
                h.append(f"<p class='muted'>评分（供需/估值/阶段）：{radar.get('supply_demand','-')}/{radar.get('valuation','-')}/{radar.get('stage','-')}　总分 {radar.get('total','-')}　平均 {radar.get('average','-')}</p>")
            h.append(f"<div class='thesis'><b>🎯 核心观点：</b><br>{esc(t.get('core_thesis',''))}</div>")
            inst = t.get("institutions") or {}
            if inst.get("list"):
                names = "、".join(f"{x.get('name','')}({x.get('report_id','')})" for x in inst["list"])
                h.append(f"<p class='muted'>🏛 机构覆盖（{inst.get('count','')}家）：{esc(names)}</p>")
            if t.get("evidence_table"):
                h.append("<table><tr><th>指标</th><th>数值</th><th>来源</th></tr>")
                for x in t["evidence_table"]:
                    h.append(f"<tr><td>{esc(x.get('indicator',''))}</td><td>{esc(x.get('value',''))}</td><td>{esc(x.get('source_type',''))}</td></tr>")
                h.append("</table>")
            if t.get("logic_chain"):
                h.append(f"<p style='white-space:pre-line;font-size:13px'><b>🧩 逻辑链条：</b><br>{esc(t['logic_chain'])}</p>")
            if t.get("risk_note"):
                h.append(f"<div class='risk'>⚠️ 风险提示：{esc(t['risk_note'])}</div>")
            if t.get("falsifiability_threshold"):
                h.append(f"<div class='falsify'>🔍 可证伪阈值：{esc(t['falsifiability_threshold'])}</div>")
            h.append("</div>")

    evs = data.get("calendar_events") or []
    if evs:
        h.append("<div class='sec'><b>四、事件日历</b></div>")
        for ev in evs:
            h.append(f"<div class='card'><b>{esc(ev.get('date',''))}　{esc(ev.get('event_name',''))}</b><br><span class='muted'>影响：{esc(ev.get('impact',''))}</span>")
            if ev.get("data_verification_point"):
                h.append(f"<div class='falsify'>✔ 验证点：{esc(ev['data_verification_point'])}</div>")
            h.append("</div>")

    corrs = data.get("correlations") or []
    if corrs:
        h.append("<div class='sec'><b>五、跨主题关联图谱</b></div>")
        for c in corrs:
            h.append(f"<div class='card'><b>{esc(c.get('from',''))} → {esc(c.get('to',''))}</b> <span class='chip'>{esc(c.get('relation',''))}</span><br><span class='muted'>{esc(c.get('evidence',''))}</span></div>")

    its = data.get("insights") or []
    if its:
        h.append("<div class='sec'><b>六、关键洞察</b></div>")
        for i, it in enumerate(its, 1):
            if isinstance(it, str):
                h.append(f"<div class='card'><b>{i}.</b> {esc(it)}</div>")
            else:
                h.append(f"<div class='card'><b>{i}.</b> {esc(it.get('title',''))} —— {esc(it.get('detail',''))}</div>")

    h.append("<p class='muted'>⚠️ 本报告仅供学习研究，不构成投资建议。数据来源：东方财富等公开渠道。</p>")
    h.append("</body></html>")
    out_path.write_text("".join(h), encoding="utf-8")
    return out_path


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--date", default="latest", help="日期 YYYY-MM-DD，默认最新")
    args = parser.parse_args()

    data = load_review(args.date)
    if data is None:
        print(f"未找到 {args.date} 的复盘数据，请先运行 fetch_data.py 和 generate_report.py。")
        return 1

    date_str = (data.get("meta") or {}).get("date", "latest")
    docx_path = OUT_DIR / f"复盘报告_{date_str}.docx"
    html_path = OUT_DIR / f"复盘报告_{date_str}.html"

    if HAS_DOCX:
        build_docx(data, date_str, docx_path)
        print(f"✅ Word 已生成: {docx_path}")
    else:
        print("⏭ 跳过 Word（未安装 python-docx）")

    build_html(data, date_str, html_path)
    print(f"✅ HTML 已生成: {html_path}")
    print("提示：用浏览器打开 HTML，按 Ctrl+P 可另存为 PDF；截图可存为图片。")

    try:
        webbrowser.open(str(html_path))
    except Exception:
        pass
    return 0


if __name__ == "__main__":
    sys.exit(main())